"""High-performance tool stack for GodBot agents."""

from __future__ import annotations

import json
import re
import shlex
import subprocess
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from duckduckgo_search import DDGS

ARTIFACTS_DIR = Path("artifacts")
ARTIFACTS_DIR.mkdir(exist_ok=True)

DANGEROUS_COMMAND_MARKERS = ["rm -rf", "mkfs", "shutdown", "reboot", ":(){", "dd if=", "chmod -R 777 /"]


@dataclass
class CachedSearch:
    ts: float
    payload: list[dict[str, Any]]


_SEARCH_CACHE: dict[str, CachedSearch] = {}
_SEARCH_TTL_SECONDS = 120


def _normalize(result: dict[str, Any]) -> dict[str, Any]:
    return {
        "title": result.get("title", ""),
        "href": result.get("href", ""),
        "body": result.get("body", ""),
        "source": result.get("source", "duckduckgo"),
    }


def web_search(query: str, max_results: int = 8, use_cache: bool = True) -> list[dict[str, Any]]:
    if not query.strip():
        return [{"error": "Query cannot be empty."}]

    cache_key = f"{query.strip().lower()}::{max_results}"
    now = time.time()
    if use_cache and cache_key in _SEARCH_CACHE:
        cached = _SEARCH_CACHE[cache_key]
        if (now - cached.ts) < _SEARCH_TTL_SECONDS:
            return cached.payload

    try:
        with DDGS() as ddgs:
            data = [_normalize(x) for x in ddgs.text(query, max_results=max_results)]
    except Exception as error:
        return [{"error": f"Web search failed: {error}"}]

    _SEARCH_CACHE[cache_key] = CachedSearch(ts=now, payload=data)
    return data


def extract_leads(text: str) -> dict[str, list[str]]:
    """Extract likely emails, domains, and phone-like patterns from raw text."""
    emails = sorted(set(re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)))
    phones = sorted(set(re.findall(r"\+?\d[\d\s\-()]{7,}\d", text)))
    domains = sorted(set(re.findall(r"\b(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}\b", text)))
    return {"emails": emails[:50], "phones": phones[:50], "domains": domains[:50]}


def save_file(filename: str, content: str) -> Path:
    safe_name = Path(filename).name
    path = ARTIFACTS_DIR / safe_name
    path.write_text(content, encoding="utf-8")
    return path


def save_json(filename: str, payload: dict[str, Any]) -> Path:
    path = ARTIFACTS_DIR / Path(filename).name
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def export_txt(filename: str, content: str) -> Path:
    """Export content as plain text file. Same as save_file but with .txt enforced."""
    base = Path(filename).stem + ".txt"
    return save_file(base, content)


def export_docx(filename: str, content: str, title: str | None = None) -> Path:
    """Export content as a Word document (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise ImportError("python-docx is required for Word export. Install with: pip install python-docx")

    doc = Document()
    if title:
        doc.add_heading(title, 0)
    for block in content.replace("\r\n", "\n").split("\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            doc.add_heading(block[2:].strip(), level=0)
        elif block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=1)
        elif block.startswith("### "):
            doc.add_heading(block[4:].strip(), level=2)
        else:
            p = doc.add_paragraph(block)
            p.paragraph_format.space_after = Pt(6)
    safe_name = Path(filename).stem + ".docx"
    path = ARTIFACTS_DIR / safe_name
    doc.save(str(path))
    return path


def export_pdf(filename: str, content: str, title: str | None = None) -> Path:
    """Export content as PDF with UTF-8 and multi-page support."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")

    safe_name = Path(filename).stem + ".pdf"
    path = ARTIFACTS_DIR / safe_name
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=40)
    styles = getSampleStyleSheet()
    story: list[Any] = []
    if title:
        story.append(Paragraph(title[:300].replace("<", "&lt;").replace(">", "&gt;"), styles["Title"]))
        story.append(Spacer(1, 12))
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=11, leading=14, spaceAfter=6
    )
    for line in content.replace("\r\n", "\n").split("\n"):
        line = line.strip()
        if not line:
            continue
        safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe, body_style))
    doc.build(story)
    return path


def execute_command(command: str, simulate: bool = True, timeout: int = 30) -> dict[str, Any]:
    if any(marker in command for marker in DANGEROUS_COMMAND_MARKERS):
        return {
            "simulated": True,
            "command": command,
            "stdout": "",
            "stderr": "Blocked dangerous command pattern.",
            "returncode": 126,
        }

    if simulate:
        return {
            "simulated": True,
            "command": command,
            "stdout": "[SIMULATION] Command not executed.",
            "stderr": "",
            "returncode": 0,
        }

    try:
        out = subprocess.run(
            shlex.split(command),
            shell=False,
            check=False,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        return {
            "simulated": False,
            "command": command,
            "stdout": out.stdout,
            "stderr": out.stderr,
            "returncode": out.returncode,
        }
    except Exception as error:
        return {
            "simulated": False,
            "command": command,
            "stdout": "",
            "stderr": str(error),
            "returncode": 1,
        }
